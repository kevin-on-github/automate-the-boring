
import docx
from datetime import date

today = date.today().strftime('%a %b %d, %Y')
d = docx.Document()
d.add_heading('This is the title', 0)
p = d.add_paragraph(f'Hello today is {today}, and this is a paragraph having ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True
d.save('workwithdocx.docx')
